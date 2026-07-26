import os
import random
from docx import Document

LICENSE_DIR = r"E:\vsc\project aws\liccense"
formats = ["2D", "2D, 3D", "2D, IMAX", "2D, 3D, IMAX"]

for filename in os.listdir(LICENSE_DIR):
    if filename.endswith(".docx") and not filename.startswith("~"):
        filepath = os.path.join(LICENSE_DIR, filename)
        
        try:
            doc = Document(filepath)
            
            # Check if it already has "Định dạng"
            has_format = False
            for para in doc.paragraphs:
                if "Định dạng:" in para.text:
                    has_format = True
                    break
                    
            if not has_format:
                fmt = random.choice(formats)
                doc.add_paragraph(f"Định dạng: {fmt}")
                doc.save(filepath)
                print(f"Updated {filename} with format: {fmt}")
            else:
                print(f"Skipped {filename}, already has format.")
        except PermissionError:
            print(f"Permission denied for {filename}. Please close the file if it's open.")
